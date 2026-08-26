import os

from routes.ads import ads_bp
from routes.blog import blog_bp
from routes.sns import sns_bp
from routes.history import history_bp
from routes.feedback import feedback_bp
from database.db import init_db, get_dashboard_data
from routes.auth import auth_bp, login_required
from routes.plan import plan_bp
from routes.poster import poster_bp
from routes.ai_office import ai_office_bp
from routes.running_form import running_form_bp
from routes.running_coach import running_coach_bp
from routes.speaking_coach import speaking_coach_bp
from routes.social_publish import social_publish_bp
from routes.payment import payment_bp
from routes.services import services_bp

from i18n.translations import SUPPORTED_LANGUAGES, TRANSLATIONS, running_i18n, translate
from i18n.speaking_copy import SPEAKING_COPY
from i18n.running_coach_copy import RUNNING_COACH_COPY

# Repair legacy mojibake that was shipped in the English speaking-coach copy.
SPEAKING_COPY.setdefault("en", {}).update({
    "hero_desc": "Describe the situation and I'll turn it into words you can actually say.",
    "input_hint": "Unstructured thoughts are okay. Sungeum will sort them out.",
    "feedback_title": "How was Teacher Sungeum's coaching? 🐶",
    "yes": "Yes, that's right",
})

SPEAKING_EXTRAS = {
    "ko": {"helpful":"👍 도움이 됐어", "neutral":"🙂 보통이야", "not_helpful":"🤔 아쉬워", "feedback_placeholder":"짧은 후기를 남겨줘 (선택)", "thanks":"고마워! 다음 말도 더 잘 도와줄게 🐶✨"},
    "en": {"helpful":"👍 Helpful", "neutral":"🙂 It was okay", "not_helpful":"🤔 Needs work", "feedback_placeholder":"Leave a short note (optional)", "thanks":"Thanks! I’ll help even better next time 🐶✨"},
    "ja": {"helpful":"👍 役に立った", "neutral":"🙂 ふつう", "not_helpful":"🤔 改善が必要", "feedback_placeholder":"短い感想をどうぞ（任意）", "thanks":"ありがとう！次はもっと上手に手伝うね 🐶✨"},
    "th": {"helpful":"👍 มีประโยชน์", "neutral":"🙂 ปานกลาง", "not_helpful":"🤔 ยังไม่ดี", "feedback_placeholder":"ฝากความคิดเห็นสั้น ๆ (ไม่บังคับ)", "thanks":"ขอบคุณนะ! ครั้งหน้าจะช่วยได้ดียิ่งขึ้น 🐶✨"},
    "zh": {"helpful":"👍 有帮助", "neutral":"🙂 一般", "not_helpful":"🤔 需要改进", "feedback_placeholder":"留下简短评价（可选）", "thanks":"谢谢！下次我会帮得更好 🐶✨"},
    "es": {"helpful":"👍 Me ayudó", "neutral":"🙂 Normal", "not_helpful":"🤔 Puede mejorar", "feedback_placeholder":"Deja un comentario breve (opcional)", "thanks":"¡Gracias! La próxima vez te ayudaré mejor 🐶✨"},
}

MENU_COPY = {
    "ko": {"광고 콘텐츠":"광고 콘텐츠", "SNS":"SNS", "블로그":"블로그", "포스터":"포스터", "러닝 코치":"순금이 러닝코치", "말하기 코치":"순금이 말하기 코치", "리뷰":"리뷰", "기록":"기록", "크레딧":"크레딧"},
    "en": {"광고 콘텐츠":"Ad content", "SNS":"SNS", "블로그":"Blog", "포스터":"Poster", "러닝 코치":"Sungeum Running Coach", "말하기 코치":"Sungeum Speaking Coach", "리뷰":"Reviews", "기록":"History", "크레딧":"Credits"},
    "ja": {"광고 콘텐츠":"広告コンテンツ", "SNS":"SNS", "블로그":"ブログ", "포스터":"ポスター", "러닝 코치":"スングム ランニングコーチ", "말하기 코치":"スングム スピーキングコーチ", "리뷰":"レビュー", "기록":"履歴", "크레딧":"クレジット"},
    "th": {"광고 콘텐츠":"เนื้อหาโฆษณา", "SNS":"SNS", "블로그":"บล็อก", "포스터":"โปสเตอร์", "러닝 코치":"โค้ชวิ่งซุนกึม", "말하기 코치":"โค้ชการพูดซุนกึม", "리뷰":"รีวิว", "기록":"ประวัติ", "크레딧":"เครดิต"},
    "zh": {"광고 콘텐츠":"广告内容", "SNS":"SNS", "블로그":"博客", "포스터":"海报", "러닝 코치":"顺金跑步教练", "말하기 코치":"顺金表达教练", "리뷰":"评价", "기록":"记录", "크레딧":"积分"},
    "es": {"광고 콘텐츠":"Contenido publicitario", "SNS":"SNS", "블로그":"Blog", "포스터":"Póster", "러닝 코치":"Entrenador de running Sungeum", "말하기 코치":"Coach de expresión Sungeum", "리뷰":"Reseñas", "기록":"Historial", "크레딧":"Créditos"},
}

# Stable semantic labels used by the menu-localizer.  Keeping these separate
# from legacy template strings prevents encoding or wording differences from
# leaving a menu item untranslated.
MENU_LABELS = {
    "ko": {"ads": "광고 콘텐츠", "sns": "SNS", "blog": "블로그", "poster": "포스터", "running": "순금이 러닝코치", "speaking": "순금이 말하기 코치", "reviews": "리뷰", "history": "기록", "credits": "크레딧"},
    "en": {"ads": "Ad content", "sns": "SNS", "blog": "Blog", "poster": "Poster", "running": "Sungeum Running Coach", "speaking": "Sungeum Speaking Coach", "reviews": "Reviews", "history": "History", "credits": "Credits"},
    "ja": {"ads": "広告コンテンツ", "sns": "SNS", "blog": "ブログ", "poster": "ポスター", "running": "スングム ランニングコーチ", "speaking": "スングム スピーキングコーチ", "reviews": "レビュー", "history": "履歴", "credits": "クレジット"},
    "th": {"ads": "เนื้อหาโฆษณา", "sns": "SNS", "blog": "บล็อก", "poster": "โปสเตอร์", "running": "โค้ชวิ่งซุนกึม", "speaking": "โค้ชการพูดซุนกึม", "reviews": "รีวิว", "history": "ประวัติ", "credits": "เครดิต"},
    "zh": {"ads": "广告内容", "sns": "SNS", "blog": "博客", "poster": "海报", "running": "顺金跑步教练", "speaking": "顺金口语教练", "reviews": "评价", "history": "记录", "credits": "积分"},
    "es": {"ads": "Contenido publicitario", "sns": "SNS", "blog": "Blog", "poster": "Póster", "running": "Entrenador de running Sungeum", "speaking": "Coach de expresión Sungeum", "reviews": "Reseñas", "history": "Historial", "credits": "Créditos"},
}
for _lang, _label in {
    "ko": "대시보드", "en": "Dashboard", "ja": "ダッシュボード",
    "th": "แดชบอร์ด", "zh": "仪表板", "es": "Panel"
}.items():
    MENU_LABELS[_lang]["dashboard"] = _label
from routes.admin import admin_bp
from routes.credits import credits_bp
from database.users import (
    init_users_table,
    set_user_admin_by_email,
)

from flask import (
    Flask,
    abort,
    render_template,
    request,
    send_file,
    send_from_directory,
    redirect,
    session
)
from werkzeug.middleware.proxy_fix import ProxyFix
from docx import Document
from docx.shared import Inches
from documents.pdf import (
    create_pdf,
    create_blog_pdf,
    create_sns_pdf,
    PDF_PATH
)

from ai.blog import make_blog
from ai.sns import make_sns

app = Flask(__name__)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1)

app.config["SECRET_KEY"] = os.environ.get(
    "SECRET_KEY",
    "project-freedom-ai-dev-secret-change-me"
)

app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["SESSION_COOKIE_SECURE"] = os.environ.get("FLASK_ENV") == "production" or bool(os.environ.get("RENDER"))
app.config["PERMANENT_SESSION_LIFETIME"] = 60 * 60 * 24 * 14


@app.before_request
def protect_same_origin_writes():
    """Reject browser write requests sent from a different site.

    This protects every existing POST endpoint without requiring each form and
    fetch call to maintain a separate token. Non-browser clients without an
    Origin header remain compatible; browser-provided cross-site origins do not.
    """
    if request.method not in {"POST", "PUT", "PATCH", "DELETE"}:
        return None
    origin = request.headers.get("Origin", "").rstrip("/")
    trusted_origins = {
        request.host_url.rstrip("/"),
        "https://projectfreedom-ai.com",
        "https://www.projectfreedom-ai.com",
        "https://project-freedom-ai.onrender.com",
    }
    configured_origins = os.environ.get("TRUSTED_ORIGINS", "")
    trusted_origins.update(
        item.strip().rstrip("/")
        for item in configured_origins.split(",")
        if item.strip()
    )
    if origin and origin not in trusted_origins:
        abort(403)
    return None


@app.after_request
def add_security_headers(response):
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("X-Frame-Options", "SAMEORIGIN")
    response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    # GPS route recommendations need the browser's location prompt. Keep
    # camera/microphone blocked by default, but allow geolocation for this
    # same-origin site so users can explicitly grant permission.
    response.headers.setdefault("Permissions-Policy", "camera=(), microphone=(), geolocation=(self)")
    if request.is_secure:
        response.headers.setdefault("Strict-Transport-Security", "max-age=31536000; includeSubDomains")
    return response

init_db()
init_users_table()

# Render/production 환경변수에 등록된 이메일을 관리자 계정으로 자동 지정
admin_email = os.environ.get("ADMIN_EMAIL", "").strip()
if admin_email:
    set_user_admin_by_email(admin_email, True)


app.register_blueprint(ads_bp)
app.register_blueprint(blog_bp)
app.register_blueprint(sns_bp)
app.register_blueprint(history_bp)
app.register_blueprint(feedback_bp)
app.register_blueprint(auth_bp)
app.register_blueprint(plan_bp)
app.register_blueprint(poster_bp)
app.register_blueprint(ai_office_bp)
app.register_blueprint(running_form_bp)
app.register_blueprint(running_coach_bp)
app.register_blueprint(speaking_coach_bp)
app.register_blueprint(social_publish_bp)
app.register_blueprint(payment_bp)
app.register_blueprint(services_bp)


# -------------------------
# 다국어 UI V1
# -------------------------

@app.context_processor
def inject_i18n():
    language = session.get("language", "ko")

    if language not in SUPPORTED_LANGUAGES:
        language = "ko"

    # Build source-text pairs so legacy hard-coded labels are localized too.
    # Existing templates already use the canonical translation values in many
    # places; this closes the remaining static text gaps without touching
    # generation or analysis logic.
    translation_pairs = []
    for key in set().union(*(TRANSLATIONS[lang].keys() for lang in SUPPORTED_LANGUAGES)):
        values = {TRANSLATIONS[lang].get(key) for lang in SUPPORTED_LANGUAGES}
        target = TRANSLATIONS[language].get(key)
        if target:
            for source in values:
                # Never translate bare numeric UI tokens.  A progress step
                # such as ``<span>3</span>`` was being mapped to the Korean
                # package label ``3개`` by the legacy fallback translator.
                if str(source).strip().isdigit() or str(target).strip().isdigit():
                    continue
                if source and source != target:
                    translation_pairs.append({"source": source, "target": target})
    for key, source in SPEAKING_COPY["ko"].items():
        target = SPEAKING_COPY.get(language, SPEAKING_COPY["ko"]).get(key)
        if source and target and source != target:
            translation_pairs.append({"source": source, "target": target})
    for source in MENU_COPY["ko"]:
        target = MENU_COPY.get(language, MENU_COPY["ko"]).get(source)
        if target and source != target:
            translation_pairs.append({"source": source, "target": target})
    for source, target_key in (("순금이 러닝코치", "러닝 코치"), ("순금이 말하기 코치", "말하기 코치")):
        target = MENU_COPY.get(language, MENU_COPY["ko"]).get(target_key)
        if target and source != target:
            translation_pairs.append({"source": source, "target": target})

    return {
        "current_language": language,
        "supported_languages": SUPPORTED_LANGUAGES,
        "t": lambda key: translate(key, language),
        "running_i18n": running_i18n(language),
        "translation_pairs": translation_pairs,
        "speaking_i18n": {**SPEAKING_COPY.get(language, SPEAKING_COPY["ko"]), **SPEAKING_EXTRAS.get(language, SPEAKING_EXTRAS["ko"])},
        "menu_i18n": MENU_LABELS.get(language, MENU_LABELS["ko"]),
        "running_coach_i18n": RUNNING_COACH_COPY.get(language, RUNNING_COACH_COPY["ko"]),
    }


@app.route("/language/<language_code>")
def set_language(language_code):
    if language_code in SUPPORTED_LANGUAGES:
        session["language"] = language_code

    next_url = request.args.get("next", "").strip()

    # 외부 URL 오픈 리다이렉트 방지: 내부 경로만 허용
    if not next_url.startswith("/") or next_url.startswith("//"):
        next_url = request.referrer or "/"

        # referrer가 외부 주소일 수 있으므로 최종적으로 안전하게 홈 사용
        if next_url.startswith("http://") or next_url.startswith("https://"):
            from urllib.parse import urlparse
            parsed = urlparse(next_url)
            next_url = parsed.path or "/"
            if parsed.query:
                next_url += "?" + parsed.query

    return redirect(next_url)



app.register_blueprint(admin_bp)
app.register_blueprint(credits_bp)

WORD_PATH = "downloads/advertisement.docx"
BLOG_WORD_PATH = "downloads/blog.docx"
SNS_WORD_PATH = "downloads/sns.docx"

# -------------------------
# 데이터베이스
# -------------------------

@app.route("/")
def landing():
    return render_template("landing.html")


@app.route("/sungeum-walk/")
def sungeum_walk():
    return send_from_directory("prototypes/sungeum-walk", "index.html")


@app.route("/sungeum-walk/<path:filename>")
def sungeum_walk_asset(filename):
    return send_from_directory("prototypes/sungeum-walk", filename)


@app.route("/terms")
def terms():
    return render_template("policy.html", policy="terms")


@app.route("/privacy")
def privacy():
    return render_template("policy.html", policy="privacy")


@app.route("/refund-policy")
def refund_policy():
    return render_template("policy.html", policy="refund")


@app.route("/dashboard")
@login_required
def dashboard():
    dashboard_data = get_dashboard_data(session["user_id"])
    return render_template("dashboard.html", dashboard=dashboard_data)


@app.route("/content")
@login_required
def content_hub():
    """Give every content format a clear, single-step entry point."""
    return render_template("content_hub.html")

# -------------------------
# Word 문서
# -------------------------

def create_word(result, image_path=""):
    os.makedirs("downloads", exist_ok=True)

    document = Document()

    # 제목
    document.add_heading(
        "Project Freedom AI",
        level=1
    )

    # AI 이미지
    if image_path and os.path.exists(image_path):
        document.add_picture(
            image_path,
            width=Inches(5.5)
        )

    # 광고 문구
    document.add_paragraph(result)

    # 저장
    document.save(WORD_PATH)

    return WORD_PATH


def create_blog_word(result, image_path=""):
    os.makedirs("downloads", exist_ok=True)

    document = Document()

    document.add_heading(
        "Project Freedom AI - Blog",
        level=1
    )

    if image_path and os.path.exists(image_path):
        document.add_picture(
            image_path,
            width=Inches(5.5)
        )

    document.add_paragraph(result)

    document.save(BLOG_WORD_PATH)

    return BLOG_WORD_PATH


def create_sns_word(result, image_path=""):
    os.makedirs("downloads", exist_ok=True)

    document = Document()

    document.add_heading(
        "Project Freedom AI - SNS",
        level=1
    )

    if image_path and os.path.exists(image_path):
        document.add_picture(
            image_path,
            width=Inches(5.5)
        )

    document.add_paragraph(result)

    document.save(SNS_WORD_PATH)

    return SNS_WORD_PATH


# -------------------------
# 광고 생성 페이지
# -------------------------

def _legacy_home_disabled():
    result = ""
    image_url = ""
    business = ""
    company = ""
    style = ""

    if request.method == "POST":
        business = request.form.get("business", "").strip()
        company = request.form.get("company", "").strip()
        style = request.form.get("style", "").strip()

        if business and company and style:

            # 1. 광고 문구 생성
            result = make_ads(
                business,
                company,
                style,
                language=session.get("language", "ko")
            )

            # 2. 이미지 생성용 프롬프트
            image_prompt = f"""
            {company}의 광고용 이미지.

            업종: {business}
            브랜드 분위기: {style}

            전문적인 SNS 광고 사진 스타일.
            고급스럽고 자연스러운 실제 사진 느낌.
            깔끔한 구성.
            이미지 안에는 글자를 넣지 말 것.
            """

            # 3. AI 이미지 생성
            image_path = make_image(image_prompt)

            # Windows 경로 → 브라우저에서 사용할 URL
            image_url = "/" + image_path.replace("\\", "/")

            # 4. 히스토리 저장
            save_history(
                business,
                company,
                style,
                result,
                image_url
            )

            create_word(result)


            # Word 생성
            create_word(
                result,
                image_path
            )

            # PDF 생성
            create_pdf(
                result,
                image_path
            )

    # 중요!!
    # 이 return은 if request.method == "POST" 안에 들어가면 안 됨
    return render_template(
        "index.html",
        result=result,
        image_url=image_url,
        business=business,
        company=company,
        style=style
    )


# -------------------------
# 블로그 생성 페이지
# -------------------------

# -------------------------
# 생성 기록 페이지
# -------------------------

# -------------------------
# Word 다운로드
# -------------------------

@app.route("/download")
def download():
    if not os.path.exists(WORD_PATH):
        return "먼저 광고를 생성해 주세요.", 404

    return send_file(
        WORD_PATH,
        as_attachment=True,
        download_name="advertisement.docx"
    )

@app.route("/download/pdf")
def download_pdf():
    if not os.path.exists(PDF_PATH):
        return "먼저 광고를 생성해 주세요.", 404

    return send_file(
        PDF_PATH,
        as_attachment=True,
        download_name="advertisement.pdf"
    )


# -------------------------
# SNS 생성 페이지
# -------------------------


@app.route("/blog/download/word")
def download_blog_word():
    if not os.path.exists(BLOG_WORD_PATH):
        return "먼저 블로그 글을 생성해 주세요.", 404

    return send_file(
        BLOG_WORD_PATH,
        as_attachment=True,
        download_name="blog.docx"
    )


@app.route("/blog/download/pdf")
def download_blog_pdf():
    blog_pdf_path = "downloads/blog.pdf"

    if not os.path.exists(blog_pdf_path):
        return "먼저 블로그 글을 생성해 주세요.", 404

    return send_file(
        blog_pdf_path,
        as_attachment=True,
        download_name="blog.pdf"
    )

@app.route("/sns/download/word")
def download_sns_word():
    if not os.path.exists(SNS_WORD_PATH):
        return "먼저 SNS 글을 생성해 주세요.", 404

    return send_file(
        SNS_WORD_PATH,
        as_attachment=True,
        download_name="sns.docx"
    )


@app.route("/sns/download/pdf")
def download_sns_pdf():
    sns_pdf_path = "downloads/sns.pdf"

    if not os.path.exists(sns_pdf_path):
        return "먼저 SNS 글을 생성해 주세요.", 404

    return send_file(
        sns_pdf_path,
        as_attachment=True,
        download_name="sns.pdf"
    )


if __name__ == "__main__":
    # Listen on the local network too, so a phone on the same Wi-Fi can preview the app.
    app.run(host="0.0.0.0", port=5000, debug=True)
