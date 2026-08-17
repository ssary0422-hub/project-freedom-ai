from pathlib import Path

from docx import Document
from docx.shared import Inches


BASE_DIR = Path(__file__).resolve().parent.parent
DOWNLOAD_DIR = BASE_DIR / "downloads"
DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)

WORD_PATH = str(DOWNLOAD_DIR / "advertisement.docx")
BLOG_WORD_PATH = str(DOWNLOAD_DIR / "blog.docx")
SNS_WORD_PATH = str(DOWNLOAD_DIR / "sns.docx")


def _create_word_document(
    output_path: str,
    title: str,
    result: str,
    image_path: str = ""
) -> str:
    document = Document()

    document.add_heading(
        title,
        level=1
    )

    if image_path:
        image_file = Path(image_path)

        if image_file.exists():
            document.add_picture(
                str(image_file),
                width=Inches(5.5)
            )

    document.add_paragraph(result or "")

    document.save(str(output_path))

    return str(output_path)


def create_word(result: str, image_path: str = "", company: str = "") -> str:
    return _create_word_document(
        WORD_PATH,
        company or "광고 콘텐츠",
        result,
        image_path
    )


def create_blog_word(result: str, image_path: str = "", company: str = "") -> str:
    return _create_word_document(
        BLOG_WORD_PATH,
        f"{company} 블로그" if company else "블로그 콘텐츠",
        result,
        image_path
    )


def create_sns_word(result: str, image_path: str = "", company: str = "") -> str:
    return _create_word_document(
        SNS_WORD_PATH,
        f"{company} SNS" if company else "SNS 콘텐츠",
        result,
        image_path
    )
