from docx import Document

document = Document()

document.add_heading("Project Freedom AI", level=1)

document.add_paragraph("안녕하세요!")

document.add_paragraph("이 문서는 Python이 만든 Word 파일입니다.")

document.save("hello.docx")

print("✅ Word 파일 생성 완료!")